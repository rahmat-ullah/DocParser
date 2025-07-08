import pytest
import asyncio
from pathlib import Path
from docx import Document
from io import BytesIO

from app.parsers.docx_parser import DOCXParser

@pytest.fixture
def docx_parser():
    return DOCXParser()

@pytest.fixture
def docx_file(tmp_path):
    # Create a simple DOCX file with content
    doc = Document()
    doc.add_heading('Test Heading', level=1)
    doc.add_paragraph('This is a test paragraph.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Row 1 Col 1'
    table.cell(1, 1).text = 'Row 1 Col 2'
    
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path

@pytest.mark.asyncio
async def test_docx_parser_supports_file(docx_parser, docx_file):
    assert docx_parser.supports_file(docx_file) == True
    
@pytest.mark.asyncio
async def test_docx_parser_parse(docx_parser, docx_file):
    result = await docx_parser.parse_to_dict(docx_file)
    assert isinstance(result, dict)
    assert "metadata" in result
    assert result["metadata"]["format"] == "DOCX"
    assert "textBlocks" in result
    assert len(result["textBlocks"]) > 0
    assert "tables" in result
    assert len(result["tables"]) > 0
    
    # Check if heading was detected
    headings = [block for block in result["textBlocks"] if block["type"] == "heading"]
    assert len(headings) > 0
    
    # Check if table was extracted
    assert len(result["tables"]) >= 1
    table = result["tables"][0]
    assert "headers" in table
    assert len(table["headers"]) == 2
