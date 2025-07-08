"""
DOCX Parser using python-docx.
Extracts text, images, and other content from DOCX files.
"""

import base64
from typing import Optional, AsyncGenerator
from pathlib import Path
from docx import Document

from .base_parser import BaseParser, ParseError
from .ast_models import DocumentAST, TextBlock, ImageBlock, TableBlock, BlockType, ParseProgress


class DOCXParser(BaseParser):
    """Parser for DOCX documents using python-docx."""

    def supports_file(self, file_path: Path) -> bool:
        """Check if file is a DOCX."""
        return file_path.suffix.lower() == '.docx'

    async def parse(
        self, file_path: Path, progress_callback: Optional[AsyncGenerator[ParseProgress, None]] = None
    ) -> DocumentAST:
        """Parse DOCX document and extract content."""
        try:
            await self._emit_progress(progress_callback, "initialization", 0.0, "Opening DOCX document")

            doc = Document(file_path)
            ast = DocumentAST(metadata={"format": "DOCX"})
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                # Determine block type
                block_type = BlockType.PARAGRAPH if not paragraph.style.name.startswith('Heading') else BlockType.HEADING
                level = self._get_heading_level(paragraph.style.name) if block_type == BlockType.HEADING else None
                
                text_block = TextBlock(
                    type=block_type,
                    content=paragraph.text.strip(),
                    level=level
                )
                ast.textBlocks.append(text_block)
                await self._emit_progress(
                    progress_callback, "parsing_paragraphs", i / len(doc.paragraphs), f"Parsing paragraph {i + 1}"
                )

            # Extract tables
            for table_i, table in enumerate(doc.tables):
                headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
                data_rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows[1:]
                ]
                if headers and data_rows:
                    table_block = TableBlock(
                        headers=headers,
                        rows=data_rows
                    )
                    ast.tables.append(table_block)
            
            # Extract images
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    img_part = rel.target_part
                    img_data = img_part.blob
                    image_base64 = base64.b64encode(img_data).decode()
                    image_block = ImageBlock(
                        data=image_base64,
                        format=img_part.content_type.split('/')[-1].upper()
                    )
                    ast.images.append(image_block)

            await self._emit_progress(progress_callback, "completion", 1.0, "DOCX parsing completed")
            return ast

        except Exception as e:
            raise ParseError(f"Failed to parse DOCX: {str(e)}", file_path, e)

    def _get_heading_level(self, style_name: str) -> int:
        """Determine heading level based on style name."""
        levels = {
            'Heading 1': 1,
            'Heading 2': 2,
            'Heading 3': 3,
            'Heading 4': 4,
        }
        return levels.get(style_name, 6) # Default to lowest priority heading
